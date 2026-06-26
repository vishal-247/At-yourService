import os
from langchain_core.documents import Document

def load_document(file_path):
    """
    Loads documents of type PDF, TXT, or DOCX.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        print(f"Loading TXT document: {file_path}")
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
        
    elif ext == ".docx":
        print(f"Loading DOCX document: {file_path}")
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = "\n".join(full_text)
        return [Document(page_content=text, metadata={"source": file_path})]
        
    elif ext == ".pdf":
        print(f"Loading PDF document: {file_path}")
        # We try to use pypdfium2 first for blazing-fast local loading (offline, no weights downloaded)
        try:
            import pypdfium2 as pdfium
            print("Using pypdfium2 loader (fast, local)...")
            pdf = pdfium.PdfDocument(file_path)
            documents = []
            for i in range(len(pdf)):
                page = pdf[i]
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                documents.append(Document(page_content=text, metadata={"source": file_path, "page": i + 1}))
            return documents
        except Exception as e:
            print(f"pypdfium2 failed or not available: {e}. Falling back to Docling...")
            from langchain_docling.loader import DoclingLoader
            loader = DoclingLoader(file_path=file_path)
            return loader.load()
            
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, TXT, and DOCX are supported.")

if __name__ == "__main__":
    # Test loading
    FILE_PATH = "./docs/rules.pdf"
    try:
        docs = load_document(FILE_PATH)
        print(f"Success! Loaded {len(docs)} document pages.")
        for idx, doc in enumerate(docs):
            print(f"Page {idx+1} length: {len(doc.page_content)}")
    except Exception as e:
        print(f"Error: {e}")
